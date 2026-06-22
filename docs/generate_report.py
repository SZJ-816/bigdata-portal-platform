#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成智讯项目设计报告"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"d:\workspace\bigdata-portal-platform\docs\project-report.docx"


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        set_cell_shading(cell, "4472C4")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后空行


def add_code_block(doc, code_text):
    """添加代码块（Consolas 字体 + 灰色底色段落）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 设置段落底色
    pPr = p._element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)
    for line in code_text.strip().split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_bullet(doc, text, level=0):
    """添加项目符号列表"""
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(11)


def add_body(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(22)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(11)


def set_heading_font(heading):
    """设置标题字体"""
    for run in heading.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def build_report():
    doc = Document()

    # ============ 页面设置 ============
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ============ 封面 ============
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("智讯 - 大数据实时科技资讯聚合平台")
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Zhixun Big Data Portal Platform")
    run.font.size = Pt(16)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("项 目 设 计 报 告")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026年6月")
    run.font.size = Pt(14)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_page_break()

    # ============ 目录 ============
    h = doc.add_heading("目  录", level=1)
    set_heading_font(h)

    toc_items = [
        "第一章 项目概述",
        "  1.1 项目背景",
        "  1.2 项目目标",
        "  1.3 技术选型",
        "第二章 系统架构设计",
        "  2.1 总体架构",
        "  2.2 后端多模块架构",
        "  2.3 数据流架构",
        "  2.4 部署架构",
        "第三章 数据库设计",
        "  3.1 MySQL 业务表",
        "  3.2 ClickHouse 分析表",
        "  3.3 Kudu 列式存储",
        "第四章 功能模块详细设计",
        "  4.1 用户端功能",
        "  4.2 后台管理功能",
        "  4.3 大数据功能",
        "  4.4 移动端功能",
        "第五章 核心代码实现",
        "  5.1 后端核心代码",
        "  5.2 前端核心代码",
        "  5.3 大数据核心代码",
        "第六章 系统测试",
        "  6.1 功能测试",
        "  6.2 性能测试",
        "  6.3 兼容性测试",
        "第七章 总结与展望",
        "  7.1 项目总结",
        "  7.2 未来展望",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        if item.startswith("  "):
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(item.strip())
            run.font.size = Pt(11)
        else:
            run = p.add_run(item)
            run.bold = True
            run.font.size = Pt(12)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_page_break()

    # ============ 第一章 项目概述 ============
    h = doc.add_heading("第一章 项目概述", level=1)
    set_heading_font(h)

    h = doc.add_heading("1.1 项目背景", level=2)
    set_heading_font(h)

    add_body(doc, "在信息爆炸的时代，科技资讯分散在各个网站和平台，用户需要花费大量时间在不同来源之间切换，才能获取全面的科技动态。这种碎片化的信息获取方式效率低下，且容易遗漏重要资讯。")
    add_body(doc, "因此，亟需一个聚合平台来整合多源科技资讯，实现一站式信息获取。通过大数据技术，可以实现实时采集、智能推荐和行为分析，为用户提供个性化、高效率的科技资讯服务。")
    add_body(doc, "本项目「智讯」应运而生，旨在利用大数据和人工智能技术，构建一个实时、智能的科技资讯聚合平台，解决信息分散、获取低效的痛点。")

    h = doc.add_heading("1.2 项目目标", level=2)
    set_heading_font(h)

    goals = [
        "构建科技资讯一站式聚合平台，整合多源RSS数据，提供统一的资讯浏览体验",
        "实现大数据实时采集与分析，基于Kafka + Flink构建实时数据处理管道",
        "AI智能摘要生成，利用NVIDIA NIM API为每篇新闻自动生成精炼摘要",
        "多端适配，覆盖Web端（Vue3）、Android端（Kotlin WebView Shell）和HarmonyOS端（ArkTS原生应用）",
    ]
    for g in goals:
        add_bullet(doc, g)

    h = doc.add_heading("1.3 技术选型", level=2)
    set_heading_font(h)

    add_body(doc, "本项目的技术选型如下表所示：")

    tech_headers = ["层次", "技术", "说明"]
    tech_rows = [
        ["前端", "Vue3 + Vite", "用户端SPA，响应式设计"],
        ["管理后台", "Vue3 + Element Plus", "后台管理系统，组件化开发"],
        ["后端", "Spring Boot 2.7 + Java 8", "多模块Maven架构，RESTful API"],
        ["数据库", "MySQL 8.0", "业务数据存储，事务支持"],
        ["大数据存储", "ClickHouse", "实时分析，列式存储引擎"],
        ["大数据存储", "Kudu", "列式存储，支持实时读写"],
        ["消息队列", "Kafka + Zookeeper", "行为数据流，高吞吐消息传递"],
        ["实时计算", "Flink", "流式处理，实时计算引擎"],
        ["AI", "NVIDIA NIM API", "智能摘要生成，大语言模型调用"],
        ["移动端", "Android (Kotlin)", "WebView Shell + 原生通知"],
        ["移动端", "HarmonyOS (ArkTS)", "原生应用，鸿蒙生态"],
        ["容器化", "Docker Compose", "一键部署，容器编排"],
    ]
    add_table(doc, tech_headers, tech_rows, col_widths=[3, 5, 7])

    doc.add_page_break()

    # ============ 第二章 系统架构设计 ============
    h = doc.add_heading("第二章 系统架构设计", level=1)
    set_heading_font(h)

    h = doc.add_heading("2.1 总体架构", level=2)
    set_heading_font(h)

    add_body(doc, "本系统采用四层架构设计，各层职责明确、松耦合，便于独立扩展和维护：")
    add_bullet(doc, "展示层：包括Vue3用户端SPA、Vue3 + Element Plus管理后台、Android WebView Shell和HarmonyOS ArkTS原生应用，负责用户交互与界面呈现")
    add_bullet(doc, "服务层：基于Spring Boot 2.7构建的微服务架构，提供RESTful API，包含认证授权、业务逻辑、AI服务等核心功能")
    add_bullet(doc, "数据层：MySQL存储业务数据，ClickHouse负责实时分析查询，Kudu用于列式存储归档数据，三者协同提供完整的数据存储方案")
    add_bullet(doc, "基础设施层：Kafka作为消息队列传递行为数据，Flink进行实时流式计算，Docker Compose实现容器化部署，Nginx提供反向代理与负载均衡")

    h = doc.add_heading("2.2 后端多模块架构", level=2)
    set_heading_font(h)

    add_body(doc, "后端采用Maven多模块架构，各模块职责如下：")

    module_headers = ["模块名称", "职责说明"]
    module_rows = [
        ["portal-admin-common", "公共模块，包含工具类、常量、通用实体等"],
        ["portal-admin-framework", "框架层，集成JWT认证、Spring Security、AOP日志等基础设施"],
        ["portal-admin-system", "系统管理模块，用户、角色、菜单权限管理"],
        ["portal-admin-cms", "内容运营模块，新闻文章、频道管理"],
        ["portal-admin-community", "用户社区模块，评论、收藏、互动功能"],
        ["portal-admin-ai", "AI智能服务模块，NVIDIA NIM API集成、摘要生成"],
        ["portal-admin-statistics", "数据统计模块，行为分析、PV/UV统计"],
        ["portal-admin-application", "启动模块，Spring Boot应用入口"],
    ]
    add_table(doc, module_headers, module_rows, col_widths=[5, 10])

    h = doc.add_heading("2.3 数据流架构", level=2)
    set_heading_font(h)

    add_body(doc, "系统的核心数据流如下：")
    add_bullet(doc, "RSS采集：定时任务从配置的RSS源抓取科技资讯，解析后存入MySQL")
    add_bullet(doc, "MySQL存储：新闻文章、用户信息、评论等业务数据持久化存储")
    add_bullet(doc, "Kafka消息队列：用户行为数据（浏览、点击、收藏等）实时发送至Kafka Topic")
    add_bullet(doc, "Flink流式处理：消费Kafka中的行为数据，进行实时聚合计算（PV/UV、热门统计等）")
    add_bullet(doc, "ClickHouse/Kudu存储：Flink处理结果写入ClickHouse用于即席查询，写入Kudu用于归档")
    add_bullet(doc, "统计分析：管理后台通过查询ClickHouse获取实时统计数据，生成可视化报表")

    h = doc.add_heading("2.4 部署架构", level=2)
    set_heading_font(h)

    add_body(doc, "系统采用Docker Compose一键部署方案：")
    add_bullet(doc, "Docker Compose编排所有服务容器，包括后端服务、MySQL、ClickHouse、Kudu、Kafka、Zookeeper、Flink等")
    add_bullet(doc, "Nginx作为反向代理，将前端请求转发至后端API，同时托管前端静态资源")
    add_bullet(doc, "前后端分离部署，前端静态文件由Nginx直接提供，后端API通过Nginx代理访问")
    add_bullet(doc, "支持水平扩展，关键服务可配置多实例负载均衡")

    doc.add_page_break()

    # ============ 第三章 数据库设计 ============
    h = doc.add_heading("第三章 数据库设计", level=1)
    set_heading_font(h)

    h = doc.add_heading("3.1 MySQL 业务表", level=2)
    set_heading_font(h)

    add_body(doc, "MySQL作为主要业务数据库，存储用户、权限、内容等核心业务数据。主要表结构如下：")

    # sys_user
    p = doc.add_paragraph()
    run = p.add_run("sys_user（用户表）")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    user_headers = ["字段名", "类型", "说明"]
    user_rows = [
        ["id", "BIGINT", "主键ID"],
        ["username", "VARCHAR(50)", "用户名"],
        ["password", "VARCHAR(100)", "密码（加密存储）"],
        ["email", "VARCHAR(100)", "邮箱"],
        ["phone", "VARCHAR(20)", "手机号"],
        ["avatar", "VARCHAR(255)", "头像URL"],
        ["status", "TINYINT", "状态（0正常 1禁用）"],
        ["role", "VARCHAR(50)", "角色标识"],
        ["create_time", "DATETIME", "创建时间"],
    ]
    add_table(doc, user_headers, user_rows, col_widths=[3, 4, 8])

    # sys_role
    p = doc.add_paragraph()
    run = p.add_run("sys_role（角色表）")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    role_headers = ["字段名", "类型", "说明"]
    role_rows = [
        ["id", "BIGINT", "主键ID"],
        ["role_name", "VARCHAR(50)", "角色名称"],
        ["role_key", "VARCHAR(50)", "角色标识"],
        ["status", "TINYINT", "状态"],
        ["create_time", "DATETIME", "创建时间"],
    ]
    add_table(doc, role_headers, role_rows, col_widths=[3, 4, 8])

    # sys_menu
    p = doc.add_paragraph()
    run = p.add_run("sys_menu（菜单权限表）")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    menu_headers = ["字段名", "类型", "说明"]
    menu_rows = [
        ["id", "BIGINT", "主键ID"],
        ["menu_name", "VARCHAR(50)", "菜单名称"],
        ["parent_id", "BIGINT", "父菜单ID"],
        ["path", "VARCHAR(200)", "路由路径"],
        ["permission", "VARCHAR(100)", "权限标识"],
        ["menu_type", "TINYINT", "类型（0目录 1菜单 2按钮）"],
    ]
    add_table(doc, menu_headers, menu_rows, col_widths=[3, 4, 8])

    # news
    p = doc.add_paragraph()
    run = p.add_run("news（新闻表）")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    news_headers = ["字段名", "类型", "说明"]
    news_rows = [
        ["id", "BIGINT", "主键ID"],
        ["title", "VARCHAR(200)", "新闻标题"],
        ["summary", "TEXT", "AI生成摘要"],
        ["content", "LONGTEXT", "新闻正文"],
        ["source", "VARCHAR(100)", "来源"],
        ["channel", "VARCHAR(50)", "频道分类"],
        ["image_url", "VARCHAR(500)", "封面图片URL"],
        ["view_count", "INT", "浏览量"],
        ["publish_time", "DATETIME", "发布时间"],
    ]
    add_table(doc, news_headers, news_rows, col_widths=[3, 4, 8])

    # 其他表简要
    other_tables = [
        ("comment（评论表）", "存储用户对新闻的评论，包含user_id、news_id、content、create_time等字段"),
        ("user_favorite（收藏表）", "存储用户收藏记录，包含user_id、news_id、create_time等字段"),
        ("user_behavior（用户行为表）", "记录用户浏览、点击等行为，包含user_id、action、target_id、create_time等字段"),
        ("user_history（浏览历史表）", "记录用户浏览历史，包含user_id、news_id、view_time等字段"),
    ]
    for table_name, desc in other_tables:
        p = doc.add_paragraph()
        run = p.add_run(table_name)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        add_body(doc, desc)

    h = doc.add_heading("3.2 ClickHouse 分析表", level=2)
    set_heading_font(h)

    add_body(doc, "ClickHouse用于实时分析场景，主要表结构如下：")

    p = doc.add_paragraph()
    run = p.add_run("user_behavior_log（行为日志表）")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    behavior_headers = ["字段名", "类型", "说明"]
    behavior_rows = [
        ["user_id", "UInt64", "用户ID"],
        ["action", "String", "行为类型（view/click/favorite/comment）"],
        ["target_id", "UInt64", "目标对象ID"],
        ["channel", "String", "频道分类"],
        ["create_time", "DateTime", "行为时间"],
    ]
    add_table(doc, behavior_headers, behavior_rows, col_widths=[3, 4, 8])

    p = doc.add_paragraph()
    run = p.add_run("pv_uv_daily（日PV/UV统计表）")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    pvuv_headers = ["字段名", "类型", "说明"]
    pvuv_rows = [
        ["stat_date", "Date", "统计日期"],
        ["pv", "UInt64", "页面浏览量"],
        ["uv", "UInt64", "独立访客数"],
        ["channel", "String", "频道分类"],
    ]
    add_table(doc, pvuv_headers, pvuv_rows, col_widths=[3, 4, 8])

    h = doc.add_heading("3.3 Kudu 列式存储", level=2)
    set_heading_font(h)

    add_body(doc, "Kudu用于新闻归档存储，支持高效的列式扫描和实时读写：")

    p = doc.add_paragraph()
    run = p.add_run("news_archive（新闻归档表）")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    archive_headers = ["字段名", "类型", "说明"]
    archive_rows = [
        ["id", "BIGINT", "主键ID"],
        ["title", "STRING", "新闻标题"],
        ["content", "STRING", "新闻正文"],
        ["source", "STRING", "来源"],
        ["channel", "STRING", "频道分类"],
        ["publish_time", "TIMESTAMP", "发布时间"],
        ["archive_time", "TIMESTAMP", "归档时间"],
    ]
    add_table(doc, archive_headers, archive_rows, col_widths=[3, 4, 8])

    doc.add_page_break()

    # ============ 第四章 功能模块详细设计 ============
    h = doc.add_heading("第四章 功能模块详细设计", level=1)
    set_heading_font(h)

    h = doc.add_heading("4.1 用户端功能", level=2)
    set_heading_font(h)

    h3 = doc.add_heading("首页", level=3)
    set_heading_font(h3)
    add_bullet(doc, "频道导航栏：展示所有频道分类，支持快速切换")
    add_bullet(doc, "新闻列表：按时间倒序展示最新资讯，支持分页加载")
    add_bullet(doc, "热门推荐：基于浏览量和互动数据推荐热门文章")

    h3 = doc.add_heading("频道页", level=3)
    set_heading_font(h3)
    add_body(doc, "支持以下频道分类浏览：AI、大数据、云计算、区块链、安全、互联网、硬件、汽车、数码、创业。每个频道独立展示该分类下的新闻列表，支持排序和筛选。")

    h3 = doc.add_heading("新闻详情", level=3)
    set_heading_font(h3)
    add_bullet(doc, "内容展示：完整的新闻正文、来源、发布时间等信息")
    add_bullet(doc, "AI摘要：一键生成新闻智能摘要，快速了解核心内容")
    add_bullet(doc, "相关推荐：基于频道和关键词推荐相关新闻")
    add_bullet(doc, "评论区：用户可发表评论，支持回复和点赞")

    h3 = doc.add_heading("搜索", level=3)
    set_heading_font(h3)
    add_bullet(doc, "全文搜索：支持按标题、内容关键词搜索新闻")
    add_bullet(doc, "搜索建议：输入时实时推荐搜索词")

    h3 = doc.add_heading("个人中心", level=3)
    set_heading_font(h3)
    add_bullet(doc, "登录注册：支持用户名密码登录和注册")
    add_bullet(doc, "收藏管理：查看和管理收藏的新闻文章")
    add_bullet(doc, "浏览历史：查看历史浏览记录，支持清除")

    h = doc.add_heading("4.2 后台管理功能", level=2)
    set_heading_font(h)

    h3 = doc.add_heading("仪表盘", level=3)
    set_heading_font(h3)
    add_bullet(doc, "用户统计：注册用户数、活跃用户数趋势")
    add_bullet(doc, "PV/UV统计：页面浏览量和独立访客数实时展示")
    add_bullet(doc, "文章统计：各频道文章数量分布")

    h3 = doc.add_heading("系统管理", level=3)
    set_heading_font(h3)
    add_bullet(doc, "用户管理：用户列表、新增、编辑、禁用等操作")
    add_bullet(doc, "角色管理：角色配置、权限分配")
    add_bullet(doc, "菜单权限：动态菜单配置，按钮级权限控制")

    h3 = doc.add_heading("内容运营", level=3)
    set_heading_font(h3)
    add_bullet(doc, "文章管理：新闻文章的增删改查、上下架操作")
    add_bullet(doc, "频道管理：频道分类的配置与管理")

    h3 = doc.add_heading("AI服务", level=3)
    set_heading_font(h3)
    add_bullet(doc, "模型配置：NVIDIA NIM API参数配置")
    add_bullet(doc, "摘要生成：手动或自动触发AI摘要生成")

    h3 = doc.add_heading("数据统计", level=3)
    set_heading_font(h3)
    add_bullet(doc, "行为分析：用户行为数据可视化分析")
    add_bullet(doc, "热门文章：基于浏览量和互动数据的热门排行")

    h = doc.add_heading("4.3 大数据功能", level=2)
    set_heading_font(h)

    add_bullet(doc, "实时采集：通过定时任务从配置的RSS源抓取科技资讯，解析标题、内容、来源等信息后存入MySQL")
    add_bullet(doc, "实时计算：Flink消费Kafka中的用户行为数据，进行窗口聚合计算，统计PV/UV和热门内容")
    add_bullet(doc, "实时分析：ClickHouse提供毫秒级即席查询能力，支持多维度的数据分析")
    add_bullet(doc, "AI摘要：调用NVIDIA NIM API，基于大语言模型为新闻自动生成精炼摘要")

    h = doc.add_heading("4.4 移动端功能", level=2)
    set_heading_font(h)

    add_bullet(doc, "Android端：采用Kotlin开发WebView Shell应用，内嵌Web端页面，集成原生推送通知能力")
    add_bullet(doc, "HarmonyOS端：采用ArkTS开发原生应用，适配鸿蒙生态，提供更流畅的交互体验")

    doc.add_page_break()

    # ============ 第五章 核心代码实现 ============
    h = doc.add_heading("第五章 核心代码实现", level=1)
    set_heading_font(h)

    h = doc.add_heading("5.1 后端核心代码", level=2)
    set_heading_font(h)

    # 统一响应体
    p = doc.add_paragraph()
    run = p.add_run("统一响应体 R<T>")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_code_block(doc, '''public class R<T> {
    private int code;
    private String msg;
    private T data;

    public static <T> R<T> ok(T data) {
        R<T> r = new R<>();
        r.setCode(200);
        r.setMsg("success");
        r.setData(data);
        return r;
    }

    public static <T> R<T> fail(String msg) {
        R<T> r = new R<>();
        r.setCode(500);
        r.setMsg(msg);
        return r;
    }
}''')

    # JWT认证过滤器
    p = doc.add_paragraph()
    run = p.add_run("JWT认证过滤器")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_code_block(doc, '''public class JwtAuthenticationFilter extends OncePerRequestFilter {
    @Override
    protected void doFilterInternal(HttpServletRequest request,
            HttpServletResponse response, FilterChain chain)
            throws ServletException, IOException {
        String token = request.getHeader("Authorization");
        if (token != null && token.startsWith("Bearer ")) {
            token = token.substring(7);
            if (JwtUtil.validateToken(token)) {
                String username = JwtUtil.getUsername(token);
                // 设置认证信息到SecurityContext
                UsernamePasswordAuthenticationToken auth =
                    new UsernamePasswordAuthenticationToken(
                        username, null, Collections.emptyList());
                SecurityContextHolder.getContext()
                    .setAuthentication(auth);
            }
        }
        chain.doFilter(request, response);
    }
}''')

    # 全局异常处理
    p = doc.add_paragraph()
    run = p.add_run("全局异常处理")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_code_block(doc, '''@RestControllerAdvice
public class GlobalExceptionHandler {
    @ExceptionHandler(Exception.class)
    public R<String> handleException(Exception e) {
        log.error("系统异常: ", e);
        return R.fail("系统异常，请联系管理员");
    }

    @ExceptionHandler(BusinessException.class)
    public R<String> handleBusinessException(BusinessException e) {
        return R.fail(e.getMessage());
    }
}''')

    # AOP操作日志
    p = doc.add_paragraph()
    run = p.add_run("AOP操作日志")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_code_block(doc, '''@Aspect
@Component
public class OperationLogAspect {
    @AfterReturning(pointcut = "@annotation(opLog)")
    public void afterReturning(JoinPoint point, OperationLog opLog) {
        // 获取当前用户
        String username = SecurityContextHolder.getContext()
            .getAuthentication().getName();
        // 记录操作日志
        SysLog log = new SysLog();
        log.setUsername(username);
        log.setOperation(opLog.value());
        log.setMethod(point.getSignature().toString());
        logService.save(log);
    }
}''')

    h = doc.add_heading("5.2 前端核心代码", level=2)
    set_heading_font(h)

    # 路由守卫
    p = doc.add_paragraph()
    run = p.add_run("路由守卫与权限控制")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_code_block(doc, '''router.beforeEach((to, from, next) => {
  const userStore = useUserStore()
  const token = userStore.token

  if (to.meta.requireAuth && !token) {
    next({ path: '/login', query: { redirect: to.fullPath } })
  } else if (to.meta.roles) {
    const userRole = userStore.role
    if (to.meta.roles.includes(userRole)) {
      next()
    } else {
      next({ path: '/403' })
    }
  } else {
    next()
  }
})''')

    # Axios拦截器
    p = doc.add_paragraph()
    run = p.add_run("Axios请求拦截器")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_code_block(doc, '''const request = axios.create({
  baseURL: '/api',
  timeout: 10000
})

request.interceptors.request.use(config => {
  const token = useUserStore().token
  if (token) {
    config.headers.Authorization = `Bearer ${token}`
  }
  return config
})

request.interceptors.response.use(
  response => response.data,
  error => {
    if (error.response?.status === 401) {
      useUserStore().logout()
      router.push('/login')
    }
    return Promise.reject(error)
  }
)''')

    # Pinia状态管理
    p = doc.add_paragraph()
    run = p.add_run("Pinia状态管理")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_code_block(doc, '''export const useUserStore = defineStore('user', () => {
  const token = ref(localStorage.getItem('token') || '')
  const userInfo = ref(null)

  async function login(formData) {
    const res = await loginApi(formData)
    token.value = res.data.token
    localStorage.setItem('token', res.data.token)
    await getUserInfo()
  }

  function logout() {
    token.value = ''
    userInfo.value = null
    localStorage.removeItem('token')
  }

  return { token, userInfo, login, logout }
})''')

    h = doc.add_heading("5.3 大数据核心代码", level=2)
    set_heading_font(h)

    # Kafka生产者
    p = doc.add_paragraph()
    run = p.add_run("Kafka行为数据生产者")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_code_block(doc, '''@Service
public class BehaviorProducer {
    @Autowired
    private KafkaTemplate<String, String> kafkaTemplate;

    public void sendBehavior(UserBehavior behavior) {
        String json = JSON.toJSONString(behavior);
        kafkaTemplate.send("user-behavior-topic",
            String.valueOf(behavior.getUserId()), json);
    }
}''')

    # Flink流式处理
    p = doc.add_paragraph()
    run = p.add_run("Flink流式处理")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_code_block(doc, '''public class BehaviorAnalysisJob {
    public static void main(String[] args) throws Exception {
        StreamExecutionEnvironment env =
            StreamExecutionEnvironment.getExecutionEnvironment();

        KafkaSource<String> source = KafkaSource.<String>builder()
            .setBootstrapServers("kafka:9092")
            .setTopics("user-behavior-topic")
            .setGroupId("flink-behavior-group")
            .setStartingOffsets(OffsetsInitializer.latest())
            .build();

        DataStream<String> stream = env.fromSource(
            source, WatermarkStrategy.noWatermarks(), "Kafka Source");

        DataStream<PvUvResult> result = stream
            .map(json -> JSON.parseObject(json, UserBehavior.class))
            .keyBy(UserBehavior::getChannel)
            .window(TumblingProcessingTimeWindows.of(Time.days(1)))
            .aggregate(new PvUvAggregateFunction());

        result.addSink(new ClickHouseSink());
        env.execute("Behavior Analysis Job");
    }
}''')

    # ClickHouse查询
    p = doc.add_paragraph()
    run = p.add_run("ClickHouse统计查询")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_code_block(doc, '''@Repository
public class ClickHouseDao {
    @Autowired
    private JdbcTemplate clickHouseJdbcTemplate;

    public List<PvUvDaily> getPvUvDaily(String startDate, String endDate) {
        String sql = "SELECT stat_date, SUM(pv) as pv, " +
            "SUM(uv) as uv FROM pv_uv_daily " +
            "WHERE stat_date BETWEEN ? AND ? " +
            "GROUP BY stat_date ORDER BY stat_date";
        return clickHouseJdbcTemplate.query(sql,
            (rs, rowNum) -> {
                PvUvDaily d = new PvUvDaily();
                d.setStatDate(rs.getString("stat_date"));
                d.setPv(rs.getLong("pv"));
                d.setUv(rs.getLong("uv"));
                return d;
            }, startDate, endDate);
    }
}''')

    doc.add_page_break()

    # ============ 第六章 系统测试 ============
    h = doc.add_heading("第六章 系统测试", level=1)
    set_heading_font(h)

    h = doc.add_heading("6.1 功能测试", level=2)
    set_heading_font(h)

    test_headers = ["测试项", "测试内容", "预期结果"]
    func_test_rows = [
        ["用户注册登录", "注册新用户并登录系统", "注册成功，登录后跳转首页"],
        ["新闻浏览与搜索", "浏览新闻列表，按关键词搜索", "列表正常展示，搜索结果准确"],
        ["后台管理CRUD", "用户、角色、文章的增删改查", "操作成功，数据正确持久化"],
        ["AI摘要生成", "对新闻文章生成AI摘要", "摘要内容准确、精炼"],
    ]
    add_table(doc, test_headers, func_test_rows, col_widths=[3.5, 5, 6.5])

    h = doc.add_heading("6.2 性能测试", level=2)
    set_heading_font(h)

    perf_headers = ["测试指标", "目标值", "说明"]
    perf_test_rows = [
        ["接口响应时间", "< 200ms", "常规CRUD接口平均响应时间"],
        ["并发处理能力", "> 500 QPS", "单机并发请求处理能力"],
        ["大数据查询性能", "< 1s", "ClickHouse即席查询响应时间"],
    ]
    add_table(doc, perf_headers, perf_test_rows, col_widths=[4, 4, 7])

    h = doc.add_heading("6.3 兼容性测试", level=2)
    set_heading_font(h)

    compat_headers = ["测试类型", "测试范围", "说明"]
    compat_test_rows = [
        ["浏览器兼容", "Chrome / Firefox / Safari / Edge", "主流浏览器最新两个大版本"],
        ["移动端适配", "iOS / Android 浏览器", "响应式布局，适配不同屏幕尺寸"],
        ["Android端", "Android 8.0+", "WebView Shell应用正常运行"],
        ["HarmonyOS端", "HarmonyOS 4.0+", "ArkTS原生应用正常运行"],
    ]
    add_table(doc, compat_headers, compat_test_rows, col_widths=[3.5, 5, 6.5])

    doc.add_page_break()

    # ============ 第七章 总结与展望 ============
    h = doc.add_heading("第七章 总结与展望", level=1)
    set_heading_font(h)

    h = doc.add_heading("7.1 项目总结", level=2)
    set_heading_font(h)

    add_body(doc, "本项目「智讯 - 大数据实时科技资讯聚合平台」已完成了全栈开发，实现了以下核心成果：")
    add_bullet(doc, "完成了科技资讯聚合平台的全栈开发，涵盖Web前端、管理后台、后端服务、移动端等完整技术栈")
    add_bullet(doc, "实现了大数据实时采集、处理、分析闭环，基于Kafka + Flink + ClickHouse构建了完整的实时数据处理管道")
    add_bullet(doc, "AI赋能内容生产，通过NVIDIA NIM API实现了新闻智能摘要生成，提升了内容消费效率")
    add_bullet(doc, "多端覆盖，Web端、Android端、HarmonyOS端三端协同，为用户提供一致的资讯体验")
    add_bullet(doc, "Docker Compose一键部署，降低了运维复杂度，提升了部署效率")

    h = doc.add_heading("7.2 未来展望", level=2)
    set_heading_font(h)

    add_body(doc, "未来计划在以下方向持续优化和扩展：")
    add_bullet(doc, "个性化推荐算法优化：引入协同过滤、深度学习等推荐算法，提升推荐精准度")
    add_bullet(doc, "更多数据源接入：扩展RSS源之外的数据采集方式，支持API对接、爬虫采集等")
    add_bullet(doc, "实时数据大屏：开发可视化数据大屏，实时展示平台运营数据和趋势")
    add_bullet(doc, "国际化支持：支持多语言界面和内容，拓展海外市场")

    # ============ 保存文件 ============
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"报告已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
